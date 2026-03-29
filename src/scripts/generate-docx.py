#!/usr/bin/env python3
"""
Academic Writer — DOCX Generator

Generates a properly formatted .docx file from article JSON data.
Handles RTL (Hebrew) text correctly, with proper punctuation and citation formatting.

Uses the Hebrew Punctuation Rules reference to ensure typography is human-like, not AI-like.

Usage:
    python3 generate-docx.py --input article.json --output path.docx

Input JSON schema:
{
    "title": str,
    "thesis": str | null,
    "sections": [{"title": str, "paragraphs": [str]}],
    "format": {
        "font": str,          # default: "David"
        "bodySize": int,       # default: 11
        "titleSize": int,      # default: 16
        "headingSize": int,    # default: 13
        "lineSpacing": float,  # default: 1.5
        "margins": float,      # default: 1.0 (inches)
        "isRtl": bool          # default: true
    },
    "totalWords": int
}
"""

import argparse
import json
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --- Threshold for showing section titles ---
SECTION_TITLE_WORD_THRESHOLD = 1500


def validate_hebrew_text(text: str) -> str:
    """
    Fix common Hebrew typographic errors that betray AI origin.
    
    Fixes:
    - Em-dashes (—) → replaced with space-hyphen-space
    - Straight quotes (") → replaced with gereshayim (״)
    - Unnecessary directional marks (RLM/LRM) → removed
    
    Args:
        text: Hebrew or mixed text string
    
    Returns:
        Cleaned text with AI typography tells removed
    """
    if not text:
        return text
    
    # 1. Replace em-dashes with space-hyphen-space
    # Em-dashes are a signature AI pattern; humans don't use them in Hebrew
    text = re.sub(r'[—\u2014]', ' - ', text)
    
    # 2. Replace straight quotes with gereshayim (Hebrew quotation marks)
    # English straight quotes are wrong in Hebrew academic writing
    text = re.sub(r'"([^"]{0,200})"', r'״\1״', text)
    
    # 3. Remove unnecessary directional marks that break rendering
    # DOCX handles bidirectionality automatically; manual marks cause issues
    
    # RLM (U+200F) before punctuation/closing brackets is usually wrong
    text = re.sub(r'\u200F+(?=[.,;:\)\]\"])', '', text)
    
    # LRM (U+200E) after opening parenthesis/bracket is usually wrong
    text = re.sub(r'(?<=[(\[])\u200E+', '', text)
    
    # Multiple consecutive directional marks should be removed
    text = re.sub(r'[\u200E\u200F]{2,}', '', text)
    
    # 4. Flag (but don't auto-fix) orphan punctuation at paragraph start
    # This would require understanding sentence structure, so we just warn
    if re.match(r'^\s*[.,;:]', text):
        # Remove the orphan punctuation
        text = re.sub(r'^\s*([.,;:])', '', text).strip()
    
    return text


def set_rtl_para(para):
    """Enable RTL layout for a paragraph via w:bidi element."""
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def set_rtl_run(run):
    """Enable RTL on a specific run via w:rtl in rPr element."""
    rPr = run._element.get_or_add_rPr()
    rtl_elem = OxmlElement("w:rtl")
    rPr.append(rtl_elem)


def add_para(doc, text, font_name, font_size, is_rtl=True,
             bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=6,
             line_spacing=1.5):
    """
    Add a paragraph with proper RTL handling for Hebrew text.
    
    Key principle: Write text naturally, let DOCX handle bidirectional rendering.
    No manual insertion of directional marks.
    
    Args:
        doc: python-docx Document
        text: Paragraph text (Hebrew or mixed)
        font_name: Font name (e.g., "David")
        font_size: Font size in points
        is_rtl: Whether text is right-to-left
        bold, italic: Text formatting
        align: Paragraph alignment
        space_before, space_after: Spacing in points
        line_spacing: Line spacing multiplier (e.g., 1.5)
    
    Returns:
        The Paragraph object
    """
    # Validate and clean text (remove AI typography tells)
    text = validate_hebrew_text(text)
    
    # Create paragraph
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = line_spacing

    if is_rtl:
        # Enable RTL context on paragraph
        set_rtl_para(para)
        
        # Add a single run with the full text
        # (no splitting into mixed-direction runs)
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        # Set complex script font
        run._element.rPr.rFonts.set(qn("w:cs"), font_name)
        # Enable RTL on run
        set_rtl_run(run)
    else:
        # LTR text (English)
        run = para.add_run(text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic

    return para


def add_page_numbers(doc):
    """Add centered page numbers in footer."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def generate_docx(data: dict, output_path: str):
    """Generate a .docx file from article data."""
    # Extract format settings with defaults
    fmt = data.get("format", {})
    font_name = fmt.get("font", "David")
    body_size = fmt.get("bodySize", 11)
    title_size = fmt.get("titleSize", 16)
    heading_size = fmt.get("headingSize", 13)
    line_spacing = fmt.get("lineSpacing", 1.5)
    margin_inches = fmt.get("margins", 1.0)
    is_rtl = fmt.get("isRtl", True)

    total_words = data.get("totalWords", 0)
    show_section_titles = total_words > SECTION_TITLE_WORD_THRESHOLD

    title = data.get("title", "Untitled")
    thesis = data.get("thesis", None)
    sections = data.get("sections", [])

    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(margin_inches)
        section.bottom_margin = Inches(margin_inches)
        section.left_margin = Inches(margin_inches)
        section.right_margin = Inches(margin_inches)

    # Title (bold, centered)
    title_align = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, title, font_name, title_size, is_rtl=is_rtl,
             bold=True, align=title_align, space_after=6,
             line_spacing=line_spacing)

    # Thesis subtitle (italic, centered) - always shown if present
    if thesis:
        add_para(doc, thesis, font_name, body_size, is_rtl=is_rtl,
                 italic=True, align=title_align, space_after=12,
                 line_spacing=line_spacing)

    # Abstract(s)
    abstract_data = data.get("abstract", None)
    if abstract_data:
        primary_abstract = abstract_data.get("primary", None)
        secondary_abstract = abstract_data.get("secondary", None)

        if primary_abstract:
            # Abstract heading
            abstract_heading = "תקציר" if is_rtl else "Abstract"
            add_para(doc, abstract_heading, font_name, heading_size,
                     is_rtl=is_rtl, bold=True,
                     align=(WD_ALIGN_PARAGRAPH.RIGHT if is_rtl
                            else WD_ALIGN_PARAGRAPH.LEFT),
                     space_before=12, space_after=6,
                     line_spacing=line_spacing)
            # Abstract body
            add_para(doc, primary_abstract, font_name, body_size,
                     is_rtl=is_rtl, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_after=6, line_spacing=line_spacing)

        if secondary_abstract:
            # Secondary abstract (e.g., English abstract for Hebrew article)
            secondary_is_rtl = not is_rtl  # opposite direction
            add_para(doc, "Abstract", font_name, heading_size,
                     is_rtl=secondary_is_rtl, bold=True,
                     align=(WD_ALIGN_PARAGRAPH.RIGHT if secondary_is_rtl
                            else WD_ALIGN_PARAGRAPH.LEFT),
                     space_before=12, space_after=6,
                     line_spacing=line_spacing)
            add_para(doc, secondary_abstract, font_name, body_size,
                     is_rtl=secondary_is_rtl,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                     space_after=12, line_spacing=line_spacing)

    # Body sections
    for sec in sections:
        # Conditional section titles
        if show_section_titles:
            heading_align = (WD_ALIGN_PARAGRAPH.RIGHT if is_rtl
                             else WD_ALIGN_PARAGRAPH.LEFT)
            add_para(doc, sec.get("title", ""), font_name, heading_size,
                     is_rtl=is_rtl, bold=True, align=heading_align,
                     space_before=12, space_after=6,
                     line_spacing=line_spacing)

        # Section paragraphs
        for para_text in sec.get("paragraphs", []):
            add_para(doc, para_text, font_name, body_size, is_rtl=is_rtl,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6,
                     line_spacing=line_spacing)

    add_page_numbers(doc)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    print(f"Words: {total_words}")
    print(f"Sections: {len(sections)}")
    print(f"Section titles: {'shown' if show_section_titles else 'hidden (under ' + str(SECTION_TITLE_WORD_THRESHOLD) + ' words)'}")


def main():
    parser = argparse.ArgumentParser(description="Generate academic article DOCX")
    parser.add_argument("--input", required=True, help="Path to article JSON file")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    generate_docx(data, args.output)


if __name__ == "__main__":
    main()
