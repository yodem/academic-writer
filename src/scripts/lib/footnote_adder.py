"""
FootnoteAdder - Add real Word footnotes to python-docx documents.

This module provides the FootnoteAdder class that enables adding native
Word footnotes to documents created with python-docx.

Usage:
    from footnote_adder import FootnoteAdder
    from docx import Document

    doc = Document("footnote_template.docx")
    footnote_adder = FootnoteAdder()

    p = doc.add_paragraph()
    p.add_run("This needs a citation")
    footnote_adder.add_footnote(p, "", "Author, Book (2024), p. 42.")

    doc.save("output.docx")
    footnote_adder.finalize_footnotes("output.docx")
"""

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import zipfile
import tempfile
import shutil
import os
import re


class FootnoteAdder:
    """Helper class to add real Word footnotes to documents."""

    def __init__(self):
        self.footnote_id = 0
        self.footnotes_to_add = []

    def add_footnote(self, paragraph, text, footnote_text):
        """Add a footnote reference to a paragraph.

        Args:
            paragraph: The python-docx paragraph object
            text: Text to add before the footnote reference (can be empty string)
            footnote_text: The footnote content

        Returns:
            The footnote run object
        """
        self.footnote_id += 1

        if text:
            paragraph.add_run(text)

        # Create the footnote reference run
        footnote_run = paragraph.add_run()
        r = footnote_run._r

        # Add run properties with FootnoteReference style
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'FootnoteReference')
        rPr.append(rStyle)
        r.insert(0, rPr)

        # Add the footnote reference element
        footnote_ref = OxmlElement('w:footnoteReference')
        footnote_ref.set(qn('w:id'), str(self.footnote_id))
        r.append(footnote_ref)

        self.footnotes_to_add.append((self.footnote_id, footnote_text))
        return footnote_run

    def finalize_footnotes(self, docx_path):
        """Add all queued footnotes to the document. Call after doc.save().

        Args:
            docx_path: Path to the saved .docx file
        """
        if not self.footnotes_to_add:
            return

        extract_dir = tempfile.mkdtemp()

        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)

            # Update footnotes.xml
            footnotes_path = os.path.join(extract_dir, "word", "footnotes.xml")
            tree = etree.parse(footnotes_path)
            root = tree.getroot()
            w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

            for fn_id, fn_text in self.footnotes_to_add:
                footnote = etree.SubElement(root, "{%s}footnote" % w_ns)
                footnote.set("{%s}id" % w_ns, str(fn_id))

                p = etree.SubElement(footnote, "{%s}p" % w_ns)
                pPr = etree.SubElement(p, "{%s}pPr" % w_ns)
                pStyle = etree.SubElement(pPr, "{%s}pStyle" % w_ns)
                pStyle.set("{%s}val" % w_ns, "FootnoteText")

                # Footnote reference mark
                r1 = etree.SubElement(p, "{%s}r" % w_ns)
                rPr1 = etree.SubElement(r1, "{%s}rPr" % w_ns)
                rStyle1 = etree.SubElement(rPr1, "{%s}rStyle" % w_ns)
                rStyle1.set("{%s}val" % w_ns, "FootnoteReference")
                etree.SubElement(r1, "{%s}footnoteRef" % w_ns)

                # Space after reference
                r2 = etree.SubElement(p, "{%s}r" % w_ns)
                t2 = etree.SubElement(r2, "{%s}t" % w_ns)
                t2.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t2.text = " "

                # Footnote text
                r3 = etree.SubElement(p, "{%s}r" % w_ns)
                t3 = etree.SubElement(r3, "{%s}t" % w_ns)
                t3.text = fn_text

            tree.write(footnotes_path, xml_declaration=True, encoding="UTF-8", standalone="yes")

            # Clean up Mac-specific content and fix XML
            self._cleanup_docx(extract_dir)

            # Repack with proper file order
            self._repack_docx(extract_dir, docx_path)

        finally:
            shutil.rmtree(extract_dir)

    def _cleanup_docx(self, extract_dir):
        """Remove Mac-specific elements and fix XML formatting."""

        for root_dir, dirs, files in os.walk(extract_dir):
            for file in files:
                if file.endswith('.xml') or file.endswith('.rels'):
                    file_path = os.path.join(root_dir, file)
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    modified = False

                    # Remove Mac namespaces
                    if 'xmlns:mo=' in content or 'xmlns:mv=' in content:
                        content = re.sub(r'\s*xmlns:mo="[^"]*"', '', content)
                        content = re.sub(r'\s*xmlns:mv="[^"]*"', '', content)
                        modified = True

                    # Fix single-quote XML declarations
                    if "<?xml version='1.0'" in content:
                        content = content.replace(
                            "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>",
                            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                        )
                        modified = True

                    if modified:
                        with open(file_path, 'w', encoding='utf-8') as f:
                            f.write(content)

        # Fix webSettings.xml
        ws_path = os.path.join(extract_dir, "word", "webSettings.xml")
        if os.path.exists(ws_path):
            with open(ws_path, 'r', encoding='utf-8') as f:
                content = f.read()
            content = re.sub(r'\s*<w:doNotSaveAsSingleFile/>', '', content)
            with open(ws_path, 'w', encoding='utf-8') as f:
                f.write(content)

        # Fix settings.xml
        settings_path = os.path.join(extract_dir, "word", "settings.xml")
        if os.path.exists(settings_path):
            with open(settings_path, 'r', encoding='utf-8') as f:
                content = f.read()
            content = re.sub(r'<w:zoom w:val="bestFit"/>', '<w:zoom w:percent="100"/>', content)
            with open(settings_path, 'w', encoding='utf-8') as f:
                f.write(content)

        # Fix docProps/app.xml
        app_path = os.path.join(extract_dir, "docProps", "app.xml")
        if os.path.exists(app_path):
            with open(app_path, 'r', encoding='utf-8') as f:
                content = f.read()
            content = content.replace('Microsoft Macintosh Word', 'Microsoft Office Word')
            content = content.replace('<Manager/>', '<Manager></Manager>')
            content = content.replace('<Company/>', '<Company></Company>')
            content = content.replace('<HyperlinkBase/>', '<HyperlinkBase></HyperlinkBase>')
            with open(app_path, 'w', encoding='utf-8') as f:
                f.write(content)

    def _repack_docx(self, extract_dir, docx_path):
        """Repack the docx with proper OOXML file order."""
        all_files = []
        for root_dir, dirs, files in os.walk(extract_dir):
            for file in files:
                file_path = os.path.join(root_dir, file)
                arcname = os.path.relpath(file_path, extract_dir).replace('\\', '/')
                all_files.append((file_path, arcname))

        # OOXML requires specific file order
        priority_order = [
            '[Content_Types].xml',
            '_rels/.rels',
            'word/_rels/document.xml.rels',
            'word/document.xml',
            'word/footnotes.xml',
            'word/endnotes.xml',
        ]

        def sort_key(item):
            try:
                return (priority_order.index(item[1]), item[1])
            except ValueError:
                return (len(priority_order), item[1])

        all_files.sort(key=sort_key)

        temp_docx = docx_path + ".tmp"
        with zipfile.ZipFile(temp_docx, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path, arcname in all_files:
                zipf.write(file_path, arcname)

        os.replace(temp_docx, docx_path)
